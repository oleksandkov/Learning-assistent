# -*- coding: utf-8 -*-
"""
Генератор звіту до лабораторної роботи №6
з UML діаграмами для веб-застосунку "Розклад занять"
"""

import os
import io
import sys

# Ensure output encoding is UTF-8
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Звіт до Лабораторної 6.docx")
PDF_PATH  = os.path.join(OUTPUT_DIR, "Звіт до Лабораторної 6.pdf")

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        element = OxmlElement(f'w:${edge}' if edge.startswith('inside') else f'w:{edge}')
        if edge in kwargs:
            for key, val in kwargs[edge].items():
                element.set(qn(f'w:{key}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_paragraph(doc, text, style='Normal', bold=False, size=None, alignment=None, space_before=None, space_after=None, color=None):
    """Add a paragraph with custom styling"""
    p = doc.add_paragraph()
    if style and style != 'Normal':
        try:
            p.style = doc.styles[style]
        except KeyError:
            pass
    
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Times New Roman'
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)
    
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    
    return p, run

def add_uml_table(doc, title, rows, col_widths=None, header_color='2D6CDF'):
    """Add a UML-style table with header styling"""
    add_styled_paragraph(doc, '', size=6, space_after=2)
    add_styled_paragraph(doc, title, bold=True, size=12, space_before=6, space_after=4)
    
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            rPr.insert(0, rFonts)
            
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, header_color)
            elif i % 2 == 0:
                set_cell_shading(cell, 'F2F6FF')
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    return table


def create_use_case_text_diagram(doc):
    """Create Use Case diagram as formatted text"""
    add_styled_paragraph(doc, '1. Діаграма варіантів використання (Use Case Diagram)', 
                         bold=True, size=13, space_before=12, space_after=6)
    
    add_styled_paragraph(doc, 
        'Діаграма варіантів використання описує функціональні вимоги до системи '
        'з точки зору користувачів (акторів). Система "Розклад занять" має два '
        'основних актори: Гість (неавторизований користувач), Звичайний Користувач '
        'та Адміністратор.', size=12, space_after=6)

    # Actors table
    actors_data = [
        ['Актор', 'Опис', 'Права'],
        ['Гість', 'Неавторизований відвідувач', 'Реєстрація, Авторизація'],
        ['Користувач', 'Авторизований користувач з роллю "user"', 'Перегляд розкладу, Фільтрація, Голосові команди (перегляд)'],
        ['Адміністратор', 'Авторизований користувач з роллю "admin"', 'Всі права Користувача + CRUD операції, Керування довідниками'],
    ]
    add_uml_table(doc, 'Таблиця 1. Актори системи', actors_data, col_widths=[3, 5, 7])

    # Use cases table
    uc_data = [
        ['ID', 'Варіант використання', 'Актор(и)', 'Опис'],
        ['UC-01', 'Реєстрація', 'Гість', 'Створення нового облікового запису з логіном та паролем'],
        ['UC-02', 'Авторизація', 'Гість', 'Вхід в систему з перевіркою облікових даних (bcrypt + JWT)'],
        ['UC-03', 'Перегляд розкладу', 'Користувач, Адмін', 'Відображення таблиці занять з групуванням за групами'],
        ['UC-04', 'Фільтрація занять', 'Користувач, Адмін', 'Пошук занять за предметом, днем, часом, аудиторією, групою, тижнем, викладачем'],
        ['UC-05', 'Додавання заняття', 'Адміністратор', 'Створення нового запису через модальну форму або голосову команду'],
        ['UC-06', 'Редагування заняття', 'Адміністратор', 'Зміна параметрів існуючого заняття (подвійний клік або кнопка)'],
        ['UC-07', 'Видалення заняття', 'Адміністратор', 'Видалення вибраного заняття з підтвердженням'],
        ['UC-08', 'Керування довідниками', 'Адміністратор', 'Додавання нових предметів, груп, викладачів до бази'],
        ['UC-09', 'Голосове введення', 'Користувач, Адмін', 'Запис голосу → STT (AssemblyAI) → текстова команда'],
        ['UC-10', 'AI-обробка команд', 'Користувач, Адмін', 'Розбір текстової команди через Google Gemini LLM для CRUD операцій'],
        ['UC-11', 'Аналіз розкладу', 'Користувач, Адмін', 'AI аналіз навчального плану з виявленням конфліктів'],
        ['UC-12', 'Вихід із системи', 'Користувач, Адмін', 'Видалення JWT токену та очищення сесії'],
    ]
    add_uml_table(doc, 'Таблиця 2. Варіанти використання', uc_data, col_widths=[1.5, 4, 3.5, 6])

    # Use Case Text Diagram
    add_styled_paragraph(doc, '', size=6)
    add_styled_paragraph(doc, 'Діаграма варіантів використання (текстове представлення):', 
                         bold=True, size=11, space_before=8, space_after=4)
    
    diagram_text = """┌─────────────────────────────────────────────────────────────────┐
│                    Система "Розклад занять"                      │
│                                                                 │
│  ┌──────────────┐  ┌──────────────────┐  ┌──────────────────┐  │
│  │  Реєстрація  │  │   Авторизація    │  │ Вихід із системи │  │
│  └──────┬───────┘  └────────┬─────────┘  └────────┬─────────┘  │
│         │                   │                     │             │
│  ┌──────┴───────────────────┴─────────────────────┴──────────┐  │
│  │                  Перегляд розкладу                         │  │
│  │  ┌─────────────────┐  ┌──────────────────────────────┐    │  │
│  │  │ Фільтрація      │  │ Голосове введення (STT)       │    │  │
│  │  │ занять           │  │  └→ AI-обробка команд (LLM)  │    │  │
│  │  └─────────────────┘  │  └→ Аналіз розкладу           │    │  │
│  │                       └──────────────────────────────────┘  │  │
│  └────────────────────────────────────────────────────────────┘  │
│                                                                 │
│  ┌────────────────── Тільки Адміністратор ──────────────────┐   │
│  │  ┌───────────────┐ ┌───────────────┐ ┌────────────────┐ │   │
│  │  │ Додавання     │ │ Редагування   │ │ Видалення      │ │   │
│  │  │ заняття       │ │ заняття       │ │ заняття        │ │   │
│  │  └───────────────┘ └───────────────┘ └────────────────┘ │   │
│  │  ┌────────────────────────────────────────────────────┐ │   │
│  │  │          Керування довідниками                      │ │   │
│  │  │  (предмети, групи, викладачі)                      │ │   │
│  │  └────────────────────────────────────────────────────┘ │   │
│  └─────────────────────────────────────────────────────────┘   │
└─────────────────────────────────────────────────────────────────┘

Актори:
 ○ Гість ────────→ Реєстрація, Авторизація
 ○ Користувач ───→ Перегляд, Фільтрація, Голосове введення, Вихід
 ○ Адміністратор → Всі функції + CRUD + Довідники"""

    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.insert(0, rFonts)


def create_class_diagram(doc):
    """Create Class Diagram as formatted tables"""
    add_styled_paragraph(doc, '2. Діаграма класів (Class Diagram)', 
                         bold=True, size=13, space_before=12, space_after=6)
    
    add_styled_paragraph(doc, 
        'Діаграма класів відображає структуру системи: серверні модулі (Express-сервер, '
        'база даних, AI-модуль) та клієнтські React-компоненти. Система побудована за '
        'архітектурою клієнт-сервер з REST API.', size=12, space_after=6)

    # Server classes
    server_classes = [
        ['Клас / Модуль', 'Тип', 'Атрибути', 'Методи'],
        ['Server (server.js)', 'Express App', 'app, port, jwtSecret, adminUsernames', 
         'auth(), requireAdmin(), createToken(), validateLesson(), normalizeTime(), normalizeDay(), fuzzyMatchTeacher(), fuzzyMatchGroup()'],
        ['Database (db.js)', 'SQLite3 модуль', 'dataDir, dbPath, db', 
         'run(), get(), all(), initDb(), seedLookups(), seedLessons(), ensureUsersStatusColumn()'],
        ['AI Module (ai.js)', 'Axios клієнти', 'geminiClient, assemblyClient, API keys', 
         'speechToText(), generateResponse(), processVoiceCommand(), parseVoiceCommand()'],
        ['App (App.jsx)', 'React Component', 'token, user, isAdmin, lookups, lessons, filters, modalOpen, editData', 
         'authSubmit(), loadLookups(), loadLessons(), saveLesson(), removeLesson(), logout()'],
        ['LessonModal', 'React Component', 'form, error, isEdit', 
         'set(), addCreateGroup(), removeCreateGroup(), submit()'],
        ['SpeechCommander', 'React Component', 'isRecording, draft, commandResult, error, status', 
         'startRecording(), stopRecording(), redoDraft(), pushCommand()'],
        ['SpeechManager', 'React Component', 'isRecording, isSpeaking, draft, detectedLanguage', 
         'startRecording(), stopRecording(), speakText(), detectLanguage(), pushCommand()'],
        ['VoiceAssistant', 'React Component', 'isListening, transcript, aiResponse, error', 
         'startListening(), stopListening()'],
        ['VoiceCommander', 'React Component', 'isRecording, transcript, commandResult', 
         'startRecording(), stopRecording(), executeCommand(), clearTranscript()'],
    ]
    add_uml_table(doc, 'Таблиця 3. Класи та модулі системи', server_classes, col_widths=[3.5, 2.5, 4, 5])

    # DB Tables
    db_tables = [
        ['Таблиця БД', 'Поля', 'Обмеження'],
        ['users', 'id (PK), username, password_hash, status, created_at', 'username UNIQUE, status IN ("user","admin")'],
        ['subjects', 'id (PK), name', 'name UNIQUE'],
        ['study_groups', 'id (PK), name', 'name UNIQUE'],
        ['teachers', 'id (PK), name', 'name UNIQUE'],
        ['lessons', 'id (PK), subject, day, lesson_time, room, group_name, week, teacher, lesson_type, created_at', 'Всі поля NOT NULL'],
    ]
    add_uml_table(doc, 'Таблиця 4. Структура бази даних (SQLite3)', db_tables, col_widths=[2.5, 7, 5])

    # Class diagram text
    add_styled_paragraph(doc, '', size=6)
    add_styled_paragraph(doc, 'Діаграма класів (текстове представлення):', 
                         bold=True, size=11, space_before=8, space_after=4)
    
    diagram_text = """┌─────────────────────────────┐     ┌──────────────────────────────┐
│       Server (Express)      │     │       Database (SQLite3)      │
├─────────────────────────────┤     ├──────────────────────────────┤
│ - app: Express              │     │ - db: sqlite3.Database       │
│ - port: number              │     │ - dbPath: string             │
│ - jwtSecret: string         │◇───→│ - dataDir: string            │
│ - adminUsernames: Set       │     ├──────────────────────────────┤
├─────────────────────────────┤     │ + run(sql, params): Promise  │
│ + auth(req, res, next)      │     │ + get(sql, params): Promise  │
│ + requireAdmin(req,res,next)│     │ + all(sql, params): Promise  │
│ + createToken(user): string │     │ + initDb(): Promise          │
│ + validateLesson(payload)   │     └──────────────────────────────┘
│ + normalizeTime(time)       │
│ + fuzzyMatchTeacher(input)  │     ┌──────────────────────────────┐
│ + fuzzyMatchGroup(input)    │     │       AI Module (ai.js)       │
└────────────┬────────────────┘     ├──────────────────────────────┤
             │                      │ - geminiClient: AxiosInstance │
             │ uses                 │ - assemblyClient: AxiosInst. │
             ▼                      │ - ASSEMBLYAI_API_KEY: string │
┌────────────────────────┐          │ - GOOGLE_GEMINI_API_KEY: str │
│   REST API Endpoints   │          ├──────────────────────────────┤
├────────────────────────┤   ◁─────│ + speechToText(buffer): str  │
│ POST /api/auth/register│          │ + generateResponse(prompt)   │
│ POST /api/auth/login   │          │ + processVoiceCommand(buf)   │
│ GET  /api/auth/me      │          │ + parseVoiceCommand(text)    │
│ GET  /api/lookups      │          └──────────────────────────────┘
│ POST /api/lookups/:e   │
│ GET  /api/lessons      │     ┌───────────────────────────────────┐
│ POST /api/lessons      │     │   React Client (App.jsx)          │
│ PUT  /api/lessons/:id  │     ├───────────────────────────────────┤
│ DELETE /api/lessons/:id│     │ - token: string                   │
│ POST /api/stt          │◁───│ - user: string                    │
│ POST /api/command      │     │ - isAdmin: boolean                │
│ POST /api/ai/chat      │     │ - lessons: Lesson[]               │
│ POST /api/ai/voice-cmd │     │ - lookups: LookupData             │
└────────────────────────┘     ├───────────────────────────────────┤
                               │ + authSubmit()                    │
                               │ + loadLookups() / loadLessons()   │
                               │ + saveLesson() / removeLesson()   │
                               └───────┬──────────┬────────────────┘
                                       │          │
                    ┌──────────────────┐│  ┌───────┴──────────────┐
                    │   LessonModal    ││  │  SpeechCommander     │
                    ├──────────────────┤│  ├──────────────────────┤
                    │ - form: FormData ││  │ - isRecording: bool  │
                    │ - error: string  ││  │ - draft: string      │
                    │ - isEdit: bool   ││  │ - commandResult: str │
                    ├──────────────────┤│  ├──────────────────────┤
                    │ + submit()       ││  │ + startRecording()   │
                    │ + addGroup()     ││  │ + stopRecording()    │
                    │ + removeGroup()  ││  │ + pushCommand()      │
                    └──────────────────┘│  └──────────────────────┘
                                       │
                          ┌────────────┴──────────────┐
                          │     SpeechManager          │
                          ├───────────────────────────┤
                          │ - isSpeaking: boolean      │
                          │ - detectedLanguage: string │
                          ├───────────────────────────┤
                          │ + speakText(text, lang)    │
                          │ + detectLanguage(text)     │
                          └───────────────────────────┘"""

    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.insert(0, rFonts)


def create_sequence_diagram(doc):
    """Create Sequence Diagram as formatted text"""
    add_styled_paragraph(doc, '3. Діаграма послідовності (Sequence Diagram)', 
                         bold=True, size=13, space_before=12, space_after=6)
    
    add_styled_paragraph(doc, 
        'Діаграма послідовності показує взаємодію компонентів системи при обробці '
        'голосової команди для додавання заняття до розкладу. '
        'Процес включає запис аудіо, розпізнавання мовлення (STT через AssemblyAI), '
        'обробку команди через AI (Google Gemini), та збереження в базу даних.', 
        size=12, space_after=6)

    # Sequence steps table
    seq_data = [
        ['Крок', 'Від', 'До', 'Повідомлення', 'Тип'],
        ['1', 'Користувач', 'SpeechCommander', 'Натискає кнопку "Record"', 'Синхронний'],
        ['2', 'SpeechCommander', 'Browser API', 'getUserMedia({audio: true})', 'Асинхронний'],
        ['3', 'Browser API', 'SpeechCommander', 'MediaStream (потік аудіо)', 'Відповідь'],
        ['4', 'SpeechCommander', 'MediaRecorder', 'start() — початок запису', 'Синхронний'],
        ['5', 'Користувач', 'SpeechCommander', 'Натискає "Stop"', 'Синхронний'],
        ['6', 'SpeechCommander', 'MediaRecorder', 'stop() — завершення запису', 'Синхронний'],
        ['7', 'MediaRecorder', 'SpeechCommander', 'audioBlob (WebM аудіо)', 'Callback'],
        ['8', 'SpeechCommander', 'Server /api/stt', 'POST {audio: base64}', 'HTTP'],
        ['9', 'Server', 'AssemblyAI API', 'Upload audio → POST /v2/upload', 'HTTP'],
        ['10', 'AssemblyAI API', 'Server', 'upload_url', 'Відповідь'],
        ['11', 'Server', 'AssemblyAI API', 'POST /v2/transcript {audio_url}', 'HTTP'],
        ['12', 'AssemblyAI API', 'Server', 'transcript.text (розпізнаний текст)', 'Polling'],
        ['13', 'Server', 'SpeechCommander', '{transcript: "Додай математику..."}', 'HTTP'],
        ['14', 'Користувач', 'SpeechCommander', 'Натискає "Ask" (Push command)', 'Синхронний'],
        ['15', 'SpeechCommander', 'Server /api/command', 'POST {prompt: draft}', 'HTTP'],
        ['16', 'Server', 'Database', 'SELECT lookups + lessons (контекст)', 'SQL'],
        ['17', 'Database', 'Server', 'Lookups + Lessons data', 'Відповідь'],
        ['18', 'Server', 'Google Gemini API', 'generateContent({prompt + context})', 'HTTP'],
        ['19', 'Google Gemini', 'Server', 'JSON: {action, parameters, ...}', 'Відповідь'],
        ['20', 'Server', 'Database', 'INSERT INTO lessons (...)', 'SQL'],
        ['21', 'Database', 'Server', 'lastID (успішне збереження)', 'Відповідь'],
        ['22', 'Server', 'SpeechCommander', '{answer: "✅ Додано заняття..."}', 'HTTP'],
        ['23', 'SpeechCommander', 'App', 'onCommandExecuted() → reload', 'Callback'],
        ['24', 'App', 'Server', 'GET /api/lessons + /api/lookups', 'HTTP'],
        ['25', 'Server', 'App', 'Оновлені дані розкладу', 'HTTP'],
    ]
    add_uml_table(doc, 'Таблиця 5. Послідовність обробки голосової команди', seq_data, col_widths=[1, 2.5, 3, 6, 2])

    # Sequence diagram text
    add_styled_paragraph(doc, '', size=6)
    add_styled_paragraph(doc, 'Діаграма послідовності (текстове представлення):', 
                         bold=True, size=11, space_before=8, space_after=4)
    
    diagram_text = """Користувач     SpeechCommander     Server          AssemblyAI      Gemini AI       Database
    │                │                 │                │               │                │
    │──── Record ───→│                 │                │               │                │
    │                │─ getUserMedia() │                │               │                │
    │                │← MediaStream ──│                │               │                │
    │                │─ start() ──────│                │               │                │
    │                │  🔴 Recording   │                │               │                │
    │──── Stop ─────→│                 │                │               │                │
    │                │─ stop() ───────│                │               │                │
    │                │← audioBlob ────│                │               │                │
    │                │─── POST /api/stt {audio} ──────→│               │                │
    │                │                 │── upload ─────→│               │                │
    │                │                 │←── upload_url ─│               │                │
    │                │                 │── transcript ─→│               │                │
    │                │                 │←── text ───────│               │                │
    │                │←── {transcript} ─│               │               │                │
    │                │  📝 Draft ready  │                │               │                │
    │──── Ask ──────→│                 │                │               │                │
    │                │── POST /api/command {prompt} ───→│               │                │
    │                │                 │── SELECT lookups, lessons ────────────────────→│
    │                │                 │←── data ──────────────────────────────────────│
    │                │                 │── generateContent(prompt+ctx) ─→│              │
    │                │                 │←── JSON {action, params} ──────│              │
    │                │                 │── INSERT INTO lessons ────────────────────────→│
    │                │                 │←── OK ────────────────────────────────────────│
    │                │←── {answer: "✅ Додано..."} ────│               │                │
    │                │─ onCommandExecuted() ──→ App refreshes data                      │
    │←── UI Updated ─│                 │                │               │                │"""

    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.insert(0, rFonts)


def create_activity_diagram(doc):
    """Create Activity Diagram"""
    add_styled_paragraph(doc, '4. Діаграма діяльності (Activity Diagram)', 
                         bold=True, size=13, space_before=12, space_after=6)
    
    add_styled_paragraph(doc, 
        'Діаграма діяльності описує потік дій при роботі з системою — від авторизації '
        'до виконання голосових команд та керування розкладом.', 
        size=12, space_after=6)

    diagram_text = """                              ● (Початок)
                              │
                              ▼
                    ┌────────────────────┐
                    │  Відкриття додатку  │
                    └─────────┬──────────┘
                              │
                              ▼
                    ┌────────────────────┐
                    │ Авторизований?     │
                    └────┬──────────┬────┘
                    Ні   │          │ Так
                         ▼          ▼
              ┌──────────────┐  ┌───────────────────┐
              │ Форма входу/ │  │ Завантаження       │
              │ реєстрації   │  │ розкладу та        │
              └──────┬───────┘  │ довідників         │
                     │          └─────────┬─────────┘
                     │  JWT token         │
                     └─────────→──────────┘
                                          │
                              ┌───────────▼───────────┐
                              │ Головна сторінка       │
                              │ (Таблиця розкладу)     │
                              └───┬──────┬──────┬─────┘
                                  │      │      │
                   ┌──────────────┘      │      └──────────────┐
                   ▼                     ▼                     ▼
        ┌──────────────────┐  ┌──────────────────┐  ┌─────────────────┐
        │  Фільтрація      │  │  Голосове         │  │  CRUD операції  │
        │  (пошук занять)  │  │  введення         │  │  (тільки адмін) │
        └──────────────────┘  └────────┬─────────┘  └────────┬────────┘
                                       │                     │
                                       ▼                     ▼
                              ┌──────────────────┐  ┌─────────────────┐
                              │ Запис аудіо      │  │ + Додавання     │
                              │ (MediaRecorder)  │  │ ✏ Редагування  │
                              └────────┬─────────┘  │ ✕ Видалення    │
                                       │            └────────┬────────┘
                                       ▼                     │
                              ┌──────────────────┐           │
                              │ STT → Транскрипт │           │
                              │ (AssemblyAI)     │           │
                              └────────┬─────────┘           │
                                       │                     │
                                       ▼                     │
                              ┌──────────────────┐           │
                              │ AI → Parse       │           │
                              │ (Google Gemini)  │           │
                              └────────┬─────────┘           │
                                       │                     │
                              ◇ Тип дії?                    │
                             ╱ │ ╲                           │
                      add  ╱  │  ╲ delete                   │
                          ╱ edit╲                            │
                         ▼   ▼   ▼                           │
                    ┌─────────────────┐                      │
                    │ Виконання       │◁────────────────────┘
                    │ операції в БД   │
                    └────────┬────────┘
                             │
                             ▼
                    ┌─────────────────┐
                    │ Оновлення UI    │
                    │ (reload data)   │
                    └────────┬────────┘
                             │
                             ▼
                             ● (Кінець циклу → повернення на головну)"""

    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(7.5)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.insert(0, rFonts)


def create_component_diagram(doc):
    """Create Component/Deployment Diagram"""
    add_styled_paragraph(doc, '5. Діаграма компонентів та розгортання (Component & Deployment Diagram)', 
                         bold=True, size=13, space_before=12, space_after=6)
    
    add_styled_paragraph(doc, 
        'Діаграма компонентів показує архітектуру розгортання системи: '
        'клієнтська частина (React через Vite), серверна частина (Express.js), '
        'база даних (SQLite3) та зовнішні сервіси (AssemblyAI, Google Gemini).', 
        size=12, space_after=6)

    comp_data = [
        ['Компонент', 'Технологія', 'Порт', 'Опис'],
        ['Frontend (Client)', 'React 18 + Vite 5', ':5173', 'SPA інтерфейс з голосовим керуванням'],
        ['Backend (Server)', 'Node.js + Express 4', ':4001', 'REST API, JWT авторизація, бізнес-логіка'],
        ['Database', 'SQLite3', 'file:schedule.db', 'Локальна БД з таблицями users, lessons, lookups'],
        ['STT Service', 'AssemblyAI API', 'HTTPS', 'Розпізнавання мовлення (universal-3-pro)'],
        ['LLM Service', 'Google Gemini 2.5 Flash', 'HTTPS', 'Обробка команд природною мовою'],
        ['Vite Proxy', 'Vite dev server', ':5173→:4001', 'Проксі /api/* запитів до Express сервера'],
    ]
    add_uml_table(doc, 'Таблиця 6. Компоненти системи', comp_data, col_widths=[3, 3.5, 2.5, 6])

    diagram_text = """┌─────────────────────────────────────────────────────────────────────────┐
│                          Клієнт (Browser)                                │
│  ┌─────────────────────────────────────────────────────────────────┐    │
│  │  React Application (Vite, port 5173)                            │    │
│  │  ┌────────────┐ ┌────────────────┐ ┌───────────────────────┐   │    │
│  │  │  App.jsx    │ │ LessonModal    │ │ SpeechCommander.jsx   │   │    │
│  │  │ (головний   │ │ (CRUD форма)   │ │ (голосове введення)   │   │    │
│  │  │  компонент) │ │                │ │ + MediaRecorder API   │   │    │
│  │  └──────┬─────┘ └──────┬─────────┘ └───────────┬───────────┘   │    │
│  │         │              │                       │               │    │
│  │         └──────────────┼───────────────────────┘               │    │
│  │                        │  fetch() / REST API                   │    │
│  └────────────────────────┼───────────────────────────────────────┘    │
│                           │ HTTP (Vite Proxy → :4001)                   │
└───────────────────────────┼─────────────────────────────────────────────┘
                            │
                            ▼
┌─────────────────────────────────────────────────────────────────────────┐
│                      Сервер (Node.js, port 4001)                        │
│  ┌──────────────────────────────────────────────────────────────────┐  │
│  │  Express.js Application (server.js)                              │  │
│  │  ┌──────────────┐  ┌──────────────┐  ┌────────────────────────┐ │  │
│  │  │ Auth Routes  │  │ CRUD Routes  │  │ AI Routes              │ │  │
│  │  │ /api/auth/*  │  │ /api/lessons │  │ /api/stt, /api/command │ │  │
│  │  │ JWT + bcrypt │  │ /api/lookups │  │ /api/ai/*              │ │  │
│  │  └──────┬───────┘  └──────┬───────┘  └──────────┬─────────────┘ │  │
│  │         │                 │                     │               │  │
│  │         ▼                 ▼                     ▼               │  │
│  │  ┌──────────────────────────────┐  ┌────────────────────────┐  │  │
│  │  │  db.js (SQLite3)             │  │  ai.js (API Clients)   │  │  │
│  │  │  schedule.db                 │  │  AssemblyAI + Gemini   │  │  │
│  │  └──────────────────────────────┘  └──────┬────────┬────────┘  │  │
│  └──────────────────────────────────────────┼────────┼────────────┘  │
└─────────────────────────────────────────────┼────────┼───────────────┘
                                              │        │
                         ┌────────────────────┘        └──────────────────┐
                         ▼                                                ▼
              ┌──────────────────────┐                      ┌──────────────────────┐
              │  AssemblyAI Cloud    │                      │  Google Cloud AI     │
              │  api.assemblyai.com  │                      │  Gemini 2.5 Flash    │
              │  ┌────────────────┐  │                      │  ┌────────────────┐  │
              │  │ Speech-to-Text │  │                      │  │ LLM Processing │  │
              │  │ (universal-3)  │  │                      │  │ (JSON output)  │  │
              │  └────────────────┘  │                      │  └────────────────┘  │
              └──────────────────────┘                      └──────────────────────┘"""

    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.insert(0, rFonts)


def create_state_diagram(doc):
    """Create State Diagram"""
    add_styled_paragraph(doc, '6. Діаграма станів (State Diagram)', 
                         bold=True, size=13, space_before=12, space_after=6)
    
    add_styled_paragraph(doc, 
        'Діаграма станів відображає можливі стани системи голосового введення '
        '(SpeechCommander) та переходи між ними.', 
        size=12, space_after=6)

    states_data = [
        ['Стан', 'Опис', 'Тригер переходу'],
        ['Idle (Ready)', 'Готовий до запису. Кнопки Record активна.', 'Натискання "Record" → Recording'],
        ['Recording', 'Запис аудіо через MediaRecorder.', 'Натискання "Stop" → Transcribing'],
        ['Transcribing', 'Відправка аудіо на сервер STT.', 'Відповідь сервера → Draft Ready / Error'],
        ['Draft Ready', 'Текст транскрипції відображено. Можна редагувати.', '"Ask" → Executing; "Clear" → Idle'],
        ['Executing', 'Відправка команди на /api/command.', 'Відповідь → Result / Error'],
        ['Result', 'Відображення результату AI обробки.', 'Новий запис → Idle'],
        ['Error', 'Відображення помилки.', 'Новий запис → Idle'],
    ]
    add_uml_table(doc, 'Таблиця 7. Стани голосового введення', states_data, col_widths=[2.5, 5, 7])

    diagram_text = """                     ● (Ініціалізація)
                     │
                     ▼
              ┌──────────────┐
              │    Idle      │◁──────────────────────────────┐
              │  (Ready)     │                               │
              └──────┬───────┘                               │
                     │ [Record]                              │
                     ▼                                       │
              ┌──────────────┐                               │
              │  Recording   │                               │
              │  🔴 Запис    │                               │
              └──────┬───────┘                               │
                     │ [Stop]                                │
                     ▼                                       │
              ┌──────────────┐                               │
              │ Transcribing │                               │
              │ ⏳ STT...    │                               │
              └───┬──────┬───┘                               │
             OK   │      │ Error                             │
                  ▼      └──────→┌──────────┐               │
              ┌──────────────┐   │  Error   │───[Clear]────→│
              │ Draft Ready  │   │  ❌      │               │
              │ 📝 Редагуван.│   └──────────┘               │
              └───┬──────┬───┘                               │
           [Ask]  │  [Clear]                                 │
                  │      └───────────────────────────────────┘
                  ▼
              ┌──────────────┐
              │  Executing   │
              │  ⚙️ AI...    │
              └───┬──────┬───┘
             OK   │      │ Error
                  ▼      └──────→┌──────────┐
              ┌──────────────┐   │  Error   │───[Clear]────→┐
              │   Result     │   │  ❌      │               │
              │   ✅ Done    │   └──────────┘               │
              └──────┬───────┘                               │
                     │ [Record / Clear]                      │
                     └───────────────────────────────────────┘"""

    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(7.5)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.insert(0, rFonts)


def create_er_diagram(doc):
    """Create ER Diagram"""
    add_styled_paragraph(doc, '7. Діаграма "Сутність-Зв\'язок" (ER Diagram)', 
                         bold=True, size=13, space_before=12, space_after=6)
    
    add_styled_paragraph(doc, 
        'ER діаграма відображає структуру бази даних SQLite3 та зв\'язки '
        'між таблицями.', 
        size=12, space_after=6)

    diagram_text = """┌──────────────────────┐         ┌──────────────────────┐
│       users          │         │      subjects         │
├──────────────────────┤         ├──────────────────────┤
│ PK id: INTEGER       │         │ PK id: INTEGER       │
│    username: TEXT     │         │    name: TEXT (UQ)    │
│    password_hash: TEXT│         └──────────┬───────────┘
│    status: TEXT       │                    │ (referenced by)
│    created_at: TEXT   │                    │
└──────────────────────┘                    │
                                            ▼
┌──────────────────────┐         ┌──────────────────────────────┐
│    study_groups       │         │          lessons              │
├──────────────────────┤         ├──────────────────────────────┤
│ PK id: INTEGER       │         │ PK id: INTEGER               │
│    name: TEXT (UQ)    │────────→│    subject: TEXT              │
└──────────────────────┘         │    day: TEXT                  │
                                 │    lesson_time: TEXT          │
┌──────────────────────┐         │    room: TEXT                 │
│      teachers         │         │    group_name: TEXT           │
├──────────────────────┤         │    week: TEXT                 │
│ PK id: INTEGER       │────────→│    teacher: TEXT              │
│    name: TEXT (UQ)    │         │    lesson_type: TEXT          │
└──────────────────────┘         │    created_at: TEXT           │
                                 └──────────────────────────────┘

Зв'язки:
  subjects.name  ←──1:N──→  lessons.subject
  study_groups.name ←──1:N──→  lessons.group_name
  teachers.name  ←──1:N──→  lessons.teacher
  users (окрема сутність для авторизації, не має прямого FK до lessons)"""

    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.insert(0, rFonts)


def generate_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # ==================== TITLE PAGE ====================
    for _ in range(3):
        doc.add_paragraph('')
    
    add_styled_paragraph(doc, 'ВІННИЦЬКИЙ НАЦІОНАЛЬНИЙ ТЕХНІЧНИЙ УНІВЕРСИТЕТ', 
                         bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, '', size=6)
    add_styled_paragraph(doc, "ФАКУЛЬТЕТ ІНФОРМАЦІЙНИХ ТЕХНОЛОГІЙ ТА КОМП'ЮТЕРНОЇ ІНЖЕНЕРІЇ", 
                         bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, '', size=6)
    add_styled_paragraph(doc, 'КАФЕДРА ПРОГРАМНОГО ЗАБЕЗПЕЧЕННЯ', 
                         bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    for _ in range(3):
        doc.add_paragraph('')
    
    add_styled_paragraph(doc, 'Звіт про виконання лабораторної роботи № 6', 
                         bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    for _ in range(4):
        doc.add_paragraph('')
    
    add_styled_paragraph(doc, 'Виконав (ла): ст. 1 курсу, групи 2ПІ-25Б ', 
                         size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_styled_paragraph(doc, 'Коваль Олександр', 
                         size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_styled_paragraph(doc, 'Перевірив: викладач Ткаченко О. М.', 
                         size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    
    for _ in range(6):
        doc.add_paragraph('')

    # ==================== CONTENT ====================
    doc.add_page_break()
    
    # Мета
    add_styled_paragraph(doc, 'Мета:', bold=True, size=12, space_after=4)
    add_styled_paragraph(doc, 
        'побудова UML діаграм для програмного забезпечення веб-застосунку '
        '"Розклад занять" з голосовим керуванням. Моделювання архітектури, '
        'поведінки та структури системи за допомогою стандартних UML нотацій.',
        size=12, space_after=8)

    # Тема
    add_styled_paragraph(doc, 'Тема:', bold=True, size=12, space_after=4)
    add_styled_paragraph(doc, 
        'Розробка UML моделей для веб-застосунку керування розкладом занять '
        'з інтеграцією штучного інтелекту (Speech-to-Text, LLM). '
        'Побудова діаграм варіантів використання, класів, послідовності, '
        'діяльності, станів, компонентів та сутність-зв\'язок.',
        size=12, space_after=8)

    # Завдання
    add_styled_paragraph(doc, 'Завдання:', bold=True, size=12, space_after=4)
    add_styled_paragraph(doc, 
        'Побудувати UML діаграми для існуючого веб-застосунку "Розклад занять", '
        'що був розроблений у лабораторних роботах 4 та 5. '
        'Діаграми повинні відображати архітектуру клієнт-серверної системи, '
        'взаємодію React-компонентів, серверних модулів (Express.js, SQLite3), '
        'та зовнішніх AI сервісів (AssemblyAI, Google Gemini).',
        size=12, space_after=8)

    # Код програми
    add_styled_paragraph(doc, 'КОД ПРОГРАМИ', bold=True, size=13, space_before=8, space_after=4)
    add_styled_paragraph(doc, 'Посилання на github автора де зберігається код програми:', size=12, space_after=2)
    p = doc.add_paragraph()
    run = p.add_run('https://github.com/oleksandkov/Learning-assistent/tree/main/PEOPLE-MACHINE/LABA5')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2D, 0x6C, 0xDF)
    run.font.name = 'Times New Roman'
    
    add_styled_paragraph(doc, '', size=6)

    # Опис стеку технологій
    add_styled_paragraph(doc, 'Стек технологій:', bold=True, size=12, space_before=8, space_after=4)
    
    tech_data = [
        ['Компонент', 'Технологія', 'Версія'],
        ['Frontend', 'React + Vite', '18.3.1 / 5.4.19'],
        ['Backend', 'Node.js + Express', '4.21.2'],
        ['База даних', 'SQLite3', '5.1.7'],
        ['Авторизація', 'JWT + bcryptjs', '9.0.2 / 3.0.2'],
        ['STT', 'AssemblyAI (universal-3-pro)', 'REST API v2'],
        ['LLM', 'Google Gemini 2.5 Flash', 'REST API v1beta'],
        ['HTTP клієнт', 'Axios', '1.7.4'],
        ['Мова', 'JavaScript (ESM)', 'ES2020+'],
    ]
    add_uml_table(doc, '', tech_data, col_widths=[3, 5, 3])

    # Structure of files
    add_styled_paragraph(doc, '', size=6)
    add_styled_paragraph(doc, 'Структура файлів проекту:', bold=True, size=12, space_before=8, space_after=4)
    
    structure = """new_laba/
├── index.html                 # Точка входу HTML
├── package.json               # Залежності та скрипти
├── vite.config.js             # Конфігурація Vite (proxy :5173→:4001)
├── .env                       # Змінні середовища (API ключі)
├── server/
│   ├── server.js              # Express сервер (1418 рядків) — REST API, auth, CRUD, AI
│   ├── db.js                  # SQLite3 модуль — ініціалізація, CRUD, seed дані
│   └── ai.js                  # AI модуль — AssemblyAI STT, Google Gemini LLM
└── src/
    ├── main.jsx               # React точка входу
    ├── App.jsx                # Головний компонент (967 рядків) — таблиця, фільтри, модалка
    ├── SpeechCommander.jsx    # Голосове введення + текстовий чат (Record→STT→Draft→Ask)
    ├── SpeechManager.jsx      # Розширене мовне керування з TTS зворотним зв'язком
    ├── VoiceAssistant.jsx     # Голосовий помічник (запис → AI обробка одним кроком)
    ├── VoiceCommander.jsx     # Голосовий командер (запис → транскрипція → виконання)
    └── style.css              # CSS стилі (450 рядків) — адаптивний дизайн"""
    
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:cs'), 'Consolas')
    rPr.insert(0, rFonts)

    # ==================== UML DIAGRAMS ====================
    doc.add_page_break()
    add_styled_paragraph(doc, 'UML ДІАГРАМИ', bold=True, size=16, 
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    
    add_styled_paragraph(doc, 
        'Нижче наведено UML діаграми, що описують архітектуру, поведінку '
        'та структуру веб-застосунку "Розклад занять" з голосовим керуванням.',
        size=12, space_after=10)

    # 1. Use Case Diagram
    create_use_case_text_diagram(doc)
    
    doc.add_page_break()
    
    # 2. Class Diagram
    create_class_diagram(doc)
    
    doc.add_page_break()
    
    # 3. Sequence Diagram
    create_sequence_diagram(doc)
    
    doc.add_page_break()
    
    # 4. Activity Diagram
    create_activity_diagram(doc)
    
    doc.add_page_break()
    
    # 5. Component Diagram
    create_component_diagram(doc)
    
    doc.add_page_break()
    
    # 6. State Diagram
    create_state_diagram(doc)
    
    doc.add_page_break()
    
    # 7. ER Diagram
    create_er_diagram(doc)

    # ==================== API ENDPOINTS ====================
    doc.add_page_break()
    add_styled_paragraph(doc, 'ОПИС REST API ЕНДПОІНТІВ', bold=True, size=14, 
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    api_data = [
        ['Метод', 'Ендпоінт', 'Авторизація', 'Опис'],
        ['GET', '/health', 'Ні', 'Перевірка стану сервера'],
        ['POST', '/api/auth/register', 'Ні', 'Реєстрація нового користувача'],
        ['POST', '/api/auth/login', 'Ні', 'Авторизація, отримання JWT'],
        ['GET', '/api/auth/me', 'JWT', 'Отримання даних поточного користувача'],
        ['GET', '/api/lookups', 'JWT', 'Завантаження довідників (предмети, групи, вчителі)'],
        ['POST', '/api/lookups/:entity', 'Admin', 'Додавання запису до довідника'],
        ['GET', '/api/lessons', 'JWT', 'Отримання списку занять (з фільтрами)'],
        ['POST', '/api/lessons', 'Admin', 'Додавання нового заняття'],
        ['PUT', '/api/lessons/:id', 'Admin', 'Редагування заняття за ID'],
        ['DELETE', '/api/lessons/:id', 'Admin', 'Видалення заняття за ID'],
        ['POST', '/api/stt', 'JWT', 'Speech-to-Text (AssemblyAI)'],
        ['POST', '/api/command', 'JWT', 'Обробка текстової команди через AI (Gemini)'],
        ['POST', '/api/ai/speech-to-text', 'JWT', 'Альтернативний STT ендпоінт'],
        ['POST', '/api/ai/chat', 'JWT', 'Чат з AI (Gemini)'],
        ['POST', '/api/ai/voice-command', 'JWT', 'Повна обробка голосової команди (STT + AI)'],
    ]
    add_uml_table(doc, 'Таблиця 8. REST API ендпоінти', api_data, col_widths=[1.5, 4, 2, 7])

    # ==================== CONCLUSION ====================
    doc.add_page_break()
    add_styled_paragraph(doc, 'Висновок:', bold=True, size=13, space_before=12, space_after=6)
    
    add_styled_paragraph(doc, 
        'У ході виконання лабораторної роботи №6 було побудовано комплект UML діаграм '
        'для веб-застосунку "Розклад занять" з голосовим керуванням, що був розроблений '
        'у попередніх лабораторних роботах.',
        size=12, space_after=4)
    
    add_styled_paragraph(doc, 
        'Було створено 7 типів UML діаграм:', 
        size=12, space_after=2)
    
    conclusions = [
        '1. Діаграма варіантів використання (Use Case) — визначено 12 варіантів використання '
        'для трьох акторів (Гість, Користувач, Адміністратор), що охоплюють авторизацію, '
        'CRUD операції над розкладом, голосове введення та AI-обробку команд.',
        
        '2. Діаграма класів (Class Diagram) — описано архітектуру з 9 основних модулів/компонентів: '
        'серверні (Server, Database, AI Module) та клієнтські (App, LessonModal, SpeechCommander, '
        'SpeechManager, VoiceAssistant, VoiceCommander), а також 5 таблиць бази даних.',
        
        '3. Діаграма послідовності (Sequence Diagram) — деталізовано 25-кроковий процес '
        'обробки голосової команди від натискання "Record" до оновлення UI з взаємодією '
        '6 компонентів (Користувач → SpeechCommander → Server → AssemblyAI → Gemini → Database).',
        
        '4. Діаграма діяльності (Activity Diagram) — відображено повний потік дій користувача '
        'в системі: від авторизації через фільтрацію та голосове введення до CRUD операцій.',
        
        '5. Діаграма компонентів та розгортання — показано архітектуру з 6 компонентів '
        'включаючи React клієнт (Vite :5173), Express сервер (:4001), SQLite3, Vite Proxy '
        'та зовнішні сервіси AssemblyAI і Google Gemini.',
        
        '6. Діаграма станів — описано 7 станів компонента SpeechCommander (Idle, Recording, '
        'Transcribing, Draft Ready, Executing, Result, Error) та переходи між ними.',
        
        '7. ER діаграма — відображено структуру бази даних з 5 таблицями (users, subjects, '
        'study_groups, teachers, lessons) та зв\'язками між ними.',
    ]
    
    for c in conclusions:
        add_styled_paragraph(doc, c, size=11, space_after=3)
    
    add_styled_paragraph(doc, '', size=6)
    add_styled_paragraph(doc, 
        'Побудовані діаграми повністю відповідають реалізованому програмному забезпеченню '
        'та демонструють архітектуру клієнт-серверного веб-застосунку з інтеграцією '
        'технологій штучного інтелекту (AssemblyAI для розпізнавання мовлення та '
        'Google Gemini 2.5 Flash для обробки команд природною мовою). '
        'Результати підтверджують відповідність системи вимогам завдання та '
        'коректність архітектурних рішень.',
        size=12, space_after=8)

    # Save DOCX
    doc.save(DOCX_PATH)
    print(f"✅ DOCX saved: {DOCX_PATH}")
    
    return DOCX_PATH


if __name__ == '__main__':
    docx_path = generate_report()
    
    # Convert to PDF
    try:
        from docx2pdf import convert
        convert(docx_path, PDF_PATH)
        print(f"✅ PDF saved: {PDF_PATH}")
    except Exception as e:
        print(f"⚠️ PDF conversion failed: {e}")
        print("You can convert DOCX to PDF manually using Word or LibreOffice")
